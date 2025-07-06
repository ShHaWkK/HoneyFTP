#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
High-Interaction FTP Honeypot (Implicit FTPS)

Fonctionnalités :
– Auto-bootstrap pip deps (Twisted, requests, service-identity, cryptography)
– Certificat TLS auto-signé dans un répertoire inscriptible  
– FTPS implicite (SSL4ServerEndpoint)  
– Auth anonymous + attacker/secret  
– Canary-files + honeytoken par session  
– FS dynamique (dirs/fichiers aléatoires) avec quota  
– RNFR/RNTO → trace dans .rename.log  
– DELE → quarantaine + log  
– MKD/RMD → alertes sur dirs “interdits”  
– Tarpitting & brute-force throttling  
– Détection Tor exit-nodes  
– SITE EXEC/CHMOD/SHELL/BOF factices  
– Virtual directory traversal (CWD ../..)  
– Logs centraux + par session
"""

import os, uuid, zipfile
import tempfile
import sys, subprocess, shutil, random, logging, smtplib
import json, atexit, base64, threading, argparse, re, mimetypes, hashlib
from twisted.internet import defer
from datetime import datetime, timedelta, timezone

# 1) Bootstrap pip deps
def ensure(pkg, imp=None):
    try:
        __import__(imp or pkg)
    except ImportError:
        subprocess.check_call([
            sys.executable,
            "-m",
            "pip",
            "install",
            "--upgrade",
            "--root-user-action=ignore",
            "--disable-pip-version-check",
            pkg,
        ])

for pkg, imp in [
    ("twisted", None),
    ("requests", None),
    ("service-identity","service_identity"),
    ("cryptography", None),
    ("pyOpenSSL", "OpenSSL"),
    ("colorama", None),
    ("rich", None),
    ("Pillow", "PIL"),
    ("openpyxl", None),
    ("python-docx", "docx"),
    ("fpdf2", "fpdf"),
]:
    ensure(pkg, imp)

import requests
from cryptography import x509
from cryptography.x509.oid import NameOID
from cryptography.hazmat.primitives import serialization, hashes
from cryptography.hazmat.primitives.asymmetric import rsa
from colorama import init as color_init, Fore, Style
from typing import Optional

# 2) Détermine le répertoire de base
script_dir = os.path.dirname(os.path.abspath(__file__))
if os.access(script_dir, os.W_OK):
    BASE = script_dir
else:
    BASE = os.path.join(tempfile.gettempdir(), "HoneyFTP")
    os.makedirs(BASE, exist_ok=True)

ROOT_DIR = os.path.join(BASE, "virtual_fs")
QUAR_DIR = os.path.join(BASE, "quarantine")
SESS_DIR = os.path.join(BASE, "sessions")
SANDBOX_DIR = os.path.join(BASE, "sandbox_reports")
LOG_FILE = os.path.join(BASE, "honeypot.log")
OP_LOG   = os.path.join(BASE, "operations.log")
PID_FILE = os.path.join(BASE, "honeypot.pid")
VERSION  = "1.1"
BANNERS = ["(vsFTPd 3.0.5)", "ProFTPD 1.3.8 Server", "FileZilla Server 0.9.60 beta", "Pure-FTPd 1.0.49"]
SERVER_START = datetime.now(timezone.utc)
EXPLICIT = os.getenv("HONEYFTP_EXPLICIT", "0") == "1"
for d in (ROOT_DIR, QUAR_DIR, SESS_DIR, SANDBOX_DIR):
    os.makedirs(d, exist_ok=True)

def _cleanup_pid():
    try:
        os.remove(PID_FILE)
    except OSError:
        pass

# 3) Logging central
color_init()

class ColorFormatter(logging.Formatter):
    COLORS = {
        logging.DEBUG: Fore.CYAN,
        logging.INFO: Fore.GREEN,
        logging.WARNING: Fore.YELLOW,
        logging.ERROR: Fore.RED,
        logging.CRITICAL: Fore.WHITE + Style.BRIGHT,
    }
    RESET = Style.RESET_ALL

    def format(self, record):
        msg = super().format(record)
        color = self.COLORS.get(record.levelno, self.RESET)
        return color + msg + self.RESET

plain_fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
color_fmt = ColorFormatter("%(asctime)s [%(levelname)s] %(message)s")

handlers = []
try:
    fh = logging.FileHandler(LOG_FILE)
    fh.setFormatter(plain_fmt)
    handlers.append(fh)
except Exception:
    pass

sh = logging.StreamHandler()
sh.setFormatter(color_fmt)
handlers.append(sh)

logging.basicConfig(level=logging.INFO, handlers=handlers)

# 4) Leurres initiaux
def create_lure_files():
    from PIL import Image
    from openpyxl import Workbook
    from docx import Document
    import io
    try:
        from fpdf import FPDF
    except Exception as e:
        FPDF = None
        logging.warning("PDF library unavailable: %s", e)

    files = {
        "passwords.txt": b"admin:admin",
        "secrets/ssh_key": b"FAKE_SSH_KEY",
        "docs/readme.txt": b"Welcome to the FTP server",
        "web/index.html": b"<html><body>Welcome</body></html>",
        "logs/syslog": b"system boot\n",
        # Liste d'employés plus complète pour rendre le leurre crédible
        "hr/employes.csv": (
            "id,name,email,dept,title,phone,hire_date\n"
            "1,Alice Smith,alice.smith@example.com,Finance,CFO,+1-555-0100,2018-01-12\n"
            "2,Bob Jones,bob.jones@example.com,IT,System Administrator,+1-555-0101,2019-03-07\n"
            "3,Charlie Ray,charlie.ray@example.com,HR,HR Manager,+1-555-0102,2017-06-21\n"
            "4,Diana Prince,diana.prince@example.com,Marketing,Lead,+1-555-0103,2020-09-30\n"
            "5,Eric Yuan,eric.yuan@example.com,Engineering,Developer,+1-555-0104,2016-11-11\n"
            "6,Francis Bean,francis.bean@example.com,Sales,Account Exec,+1-555-0105,2021-05-19\n"
            "7,Gwen Stark,gwen.stark@example.com,Support,Technician,+1-555-0106,2015-02-14\n"
            "8,Hugh Grant,hugh.grant@example.com,Finance,Accountant,+1-555-0107,2018-12-01\n"
            "9,Ivy Chen,ivy.chen@example.com,IT,DevOps,+1-555-0108,2019-07-23\n"
            "10,John Doe,john.doe@example.com,Management,CEO,+1-555-0109,2014-04-04\n"
        ).encode("utf-8"),
        "docs/README.md": b"Internal documentation\n",
        # Les identifiants incluent l'adresse du serveur SSH leurre
        "credentials.json": (
            "{\n"
            "  \"user\": \"admin\",\n"
            "  \"pass\": \"secret\",\n"
            "  \"server\": \"SSH 192.168.100.51:2224\"\n"
            "}"
        ).encode("utf-8"),
    }

    if FPDF:
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            pdf.cell(40, 10, "Manual")
            data = pdf.output(dest="S")
            if isinstance(data, str):
                data = data.encode("latin1")
            else:
                data = bytes(data)
            files["docs/manual.pdf"] = data
        except Exception as e:
            logging.warning("Failed to generate PDF lure: %s", e)

    buf = io.BytesIO()
    Image.new("RGB", (1, 1), (255, 0, 0)).save(buf, format="JPEG")
    files["images/photo.jpg"] = buf.getvalue()

    buf = io.BytesIO()
    Image.new("RGB", (1, 1), (0, 255, 0)).save(buf, format="PNG")
    files["public/images/confidentiel.png"] = buf.getvalue()

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        z.writestr("logs.txt", "log")
    files["backups/backup.zip"] = buf.getvalue()
    files["backups/logs_2025-07-04.zip"] = buf.getvalue()

    wb = Workbook()
    ws = wb.active
    ws.title = "Financials"
    ws.append(["Quarter", "Revenue", "Expenses", "Profit"])
    for q in range(1,5):
        rev = random.randint(50000, 150000)
        exp = random.randint(20000, 80000)
        ws.append([f"Q{q}", rev, exp, rev-exp])
    buf = io.BytesIO()
    wb.save(buf)
    files["finance/Financials.xlsx"] = buf.getvalue()

    doc = Document()
    doc.add_heading("Project Apollo", 0)
    doc.add_paragraph(
        "This document outlines the objectives and milestones of Project Apollo, "
        "our next generation platform."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List"
    hdr = table.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Description"
    for phase in ["Design", "Implementation", "Testing", "Deployment"]:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = f"Details for {phase.lower()} phase"
    buf = io.BytesIO()
    doc.save(buf)
    files["docs/project_apollo.docx"] = buf.getvalue()

    key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    priv = key.private_bytes(
        serialization.Encoding.PEM,
        serialization.PrivateFormat.TraditionalOpenSSL,
        serialization.NoEncryption(),
    )
    files["it/ssh_keys/id_rsa"] = priv
    files["aws_keys.pem"] = priv

    files["uploads/.keep"] = b""

    for rel, data in files.items():
        fp = os.path.join(ROOT_DIR, rel)
        os.makedirs(os.path.dirname(fp), exist_ok=True)
        if not os.path.exists(fp):
            mode = "wb" if isinstance(data, (bytes, bytearray)) else "w"
            with open(fp, mode) as f:
                f.write(data)

create_lure_files()

failed_attempts = {}
dynamic_items   = []
# Counters for SITE STATS
STATS = {
    "connections": 0,
    "logins": 0,
    "uploads": 0,
    "downloads": 0,
    "deletes": 0,
    "renames": 0,
    "ls": 0,
    "cd": 0,
}

# Persist global statistics across restarts
STATS_FILE = os.path.join(BASE, "stats_global.json")

def load_stats() -> None:
    try:
        with open(STATS_FILE) as f:
            data = json.load(f)
            STATS.update({k: int(data.get(k, v)) for k, v in STATS.items()})
    except Exception:
        pass

def save_stats() -> None:
    try:
        with open(STATS_FILE, "w") as f:
            json.dump(STATS, f)
    except Exception:
        pass

load_stats()

# Register cleanup handlers
atexit.register(_cleanup_pid)
atexit.register(save_stats)

# Suggested post-session actions to display in the menu
ACTIONS = []

def get_next_session_id() -> str:
    """Return the next incremental session ID (session1, session2, ...)."""
    max_id = 0
    if os.path.isdir(SESS_DIR):
        for name in os.listdir(SESS_DIR):
            m = re.search(r"session(\d+)", name)
            if m:
                num = int(m.group(1))
                if num > max_id:
                    max_id = num
    return f"session{max_id + 1}"

# Limite maximale d'espace disque pour le faux filesystem (env HONEYFTP_QUOTA_MB)
QUOTA_BYTES = int(os.getenv("HONEYFTP_QUOTA_MB", "10")) * 1024 * 1024
# Quota disque par session (50 Mo par défaut)
SESSION_QUOTA = int(os.getenv("HONEYFTP_SESSION_QUOTA_MB", "50")) * 1024 * 1024

# 5) Génération du certificat TLS
KEY_FILE = os.path.join(ROOT_DIR, "server.key")
CRT_FILE = os.path.join(ROOT_DIR, "server.crt")
if not (os.path.exists(KEY_FILE) and os.path.exists(CRT_FILE)):
    key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    with open(KEY_FILE, "wb") as f:
        f.write(key.private_bytes(
            serialization.Encoding.PEM,
            serialization.PrivateFormat.TraditionalOpenSSL,
            serialization.NoEncryption()))
    subj = x509.Name([
        x509.NameAttribute(NameOID.COUNTRY_NAME,           u"US"),
        x509.NameAttribute(NameOID.STATE_OR_PROVINCE_NAME, u"CA"),
        x509.NameAttribute(NameOID.LOCALITY_NAME,          u"San Francisco"),
        x509.NameAttribute(NameOID.ORGANIZATION_NAME,      u"Honeypot"),
        x509.NameAttribute(NameOID.COMMON_NAME,            u"localhost"),
    ])
    cert = (
        x509.CertificateBuilder()
           .subject_name(subj).issuer_name(subj)
           .public_key(key.public_key())
           .serial_number(x509.random_serial_number())
           .not_valid_before(datetime.now(timezone.utc))
           .not_valid_after(datetime.now(timezone.utc)+timedelta(days=365))
           .add_extension(
               x509.SubjectAlternativeName([x509.DNSName(u"localhost")]),
               critical=False
           )
           .sign(key, hashes.SHA256())
    )
    with open(CRT_FILE, "wb") as f:
        f.write(cert.public_bytes(serialization.Encoding.PEM))

# 6) Twisted & config
from twisted.cred import portal, checkers, credentials, error
from twisted.cred.checkers import AllowAnonymousAccess
from twisted.internet import endpoints, reactor, ssl, defer, error as net_error
from twisted.internet.protocol import DatagramProtocol
from twisted.protocols import ftp
from twisted.python import filepath

TOR_LIST   = "https://check.torproject.org/torbulkexitlist"
BRUTEF_THR = 5
DELAY_SEC  = 2
CANARY     = {
    "passwords.txt",
    "secrets/ssh_key",
    "credentials.json",
    "aws_keys.pem",
}
FORBID     = {"secrets"}        # supprimer un répertoire "secrets" déclenche alerta
PORT       = int(os.getenv("HONEYFTP_PORT","2121"))
KNOCK_SEQ  = [4020, 4021, 4022]
SLACK_URL  = os.getenv("SLACK_WEBHOOK")
# Configuration SMTP par variables d'environnement avec valeurs par défaut
SMTP_HOST = os.getenv("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", 587))
SMTP_USER = os.getenv("SMTP_USER", "honeycute896@gmail.com")
SMTP_PASS = os.getenv("SMTP_PASS", "jawm fmcm dmaf qkyl")
ALERT_FROM = SMTP_USER
ALERT_TO = os.getenv("ALERT_TO", "alexandreuzan75@gmail.com")
SMTP_CFG   = (
    SMTP_HOST,
    SMTP_PORT,
    SMTP_USER,
    SMTP_PASS,
    ALERT_FROM,
    ALERT_TO,
)

def alert(
    msg: str,
    *,
    ip: Optional[str] = None,
    user: Optional[str] = None,
    session: Optional[str] = None,
    log_file: Optional[str] = None,
) -> None:
    """Send an alert via Slack and/or SMTP with contextual information."""
    ts = datetime.now(timezone.utc).isoformat()
    details = [msg, "", f"Time: {ts}"]
    if ip:
        details.append(f"IP: {ip}")
    if user:
        details.append(f"User: {user}")
    if session:
        details.append(f"Session: {session}")
    if log_file:
        try:
            with open(log_file) as f:
                lines = f.read().splitlines()[-5:]
            details.append("")
            details.append("Last commands:")
            details.extend(lines)
        except Exception:
            pass
    body = "\n".join(details)

    if SLACK_URL:
        try:
            requests.post(SLACK_URL, json={"text": body}, timeout=5)
        except Exception:
            pass
    srv, port, u, pw, fr, to = SMTP_CFG
    if srv and fr and to:
        try:
            s = smtplib.SMTP(srv, port, timeout=5)
            if u:
                s.starttls()
                s.login(u, pw or "")
            mail = f"Subject:HoneyFTP Alert\nFrom:{fr}\nTo:{to}\n\n{body}"
            s.sendmail(fr, [to], mail)
            s.quit()
        except Exception:
            pass

def log_operation(msg: str):
    """Append an entry to the operations log."""
    try:
        with open(OP_LOG, "a") as f:
            ts = datetime.now(timezone.utc).isoformat()
            f.write(f"{ts} {msg}\n")
    except Exception:
        pass

def sandbox_analyze(path: str, session: str, user: str, ip: str) -> None:
    """Run basic sandbox analysis on an uploaded file."""
    report = {
        "file": os.path.basename(path),
        "size": os.path.getsize(path),
        "mtime": datetime.fromtimestamp(os.path.getmtime(path), tz=timezone.utc).isoformat(),
        "md5": None,
        "sha256": None,
        "mime": mimetypes.guess_type(path)[0] or "unknown",
        "strings": "",
    }
    try:
        data = open(path, "rb").read()
        report["md5"] = hashlib.md5(data).hexdigest()
        report["sha256"] = hashlib.sha256(data).hexdigest()
    except Exception as e:
        report["error"] = str(e)
    try:
        proc = subprocess.run(
            ["firejail", "--quiet", "--private", "--net=none", "strings", path],
            capture_output=True, text=True, timeout=10
        )
        out = proc.stdout.splitlines()[:20]
        report["strings"] = "\n".join(out)
    except Exception as e:
        report["strings"] = f"sandbox error: {e}"
    os.makedirs(SANDBOX_DIR, exist_ok=True)
    rep_path = os.path.join(SANDBOX_DIR, f"{session}_{uuid.uuid4().hex}.json")
    try:
        with open(rep_path, "w") as f:
            json.dump(report, f, indent=2)
        alert(
            f"Sandbox report for {report['file']}",
            ip=ip,
            user=user,
            session=session,
        )
    except Exception:
        pass

TOR_CACHE = {"ts": None, "ips": set()}

def is_tor_exit(ip: str) -> bool:
    """Vérifie si l'IP provient d'un noeud de sortie Tor (cache 1h)."""
    now = datetime.now(timezone.utc)
    if not TOR_CACHE["ts"] or (now - TOR_CACHE["ts"]) > timedelta(hours=1):
        try:
            r = requests.get(TOR_LIST, timeout=5)
            TOR_CACHE["ips"] = set(r.text.splitlines())
            TOR_CACHE["ts"] = now
        except Exception:
            return False
    return ip in TOR_CACHE["ips"]

def create_honeytoken(ip: str, sess: str) -> str:
    fn = f"secret_{sess}.txt"
    with open(os.path.join(ROOT_DIR, fn),"w") as f:
        f.write(f"session={sess}\nip={ip}\n")
    return fn

def create_user_lure(user: str) -> str:
    """Create a user-specific lure file."""
    fn = f"notes_{user}.txt"
    path = os.path.join(ROOT_DIR, fn)
    try:
        with open(path, "w") as f:
            f.write(f"Private notes for {user}\n")
        log_operation(f"LURE {fn}")
    except Exception:
        pass
    return fn

def validate_path(rel: str, cwd=None):
    """Validate a client-supplied path taking the current directory into account."""
    try:
        segs = ftp.toSegments(cwd or [], rel)
    except Exception:
        raise ValueError("Invalid path")
    abs_path = os.path.abspath(os.path.join(ROOT_DIR, *segs))
    if not abs_path.startswith(os.path.abspath(ROOT_DIR)):
        raise ValueError("Invalid path")
    return abs_path, "/".join(segs)

def disk_usage(path: str) -> int:
    """Retourne la taille totale (en octets) d'un répertoire."""
    total = 0
    for root, dirs, files in os.walk(path):
        for f in files:
            fp = os.path.join(root, f)
            try:
                total += os.path.getsize(fp)
            except OSError:
                pass
    return total

def finalize_session(sess: str, start: datetime, dls=0, ups=0, cds=0, rns=0, size=0):
    """Archive logs and stats for a session"""
    end = datetime.now(timezone.utc)
    stats_path = os.path.join(SESS_DIR, f"stats_{sess}.json")
    data = {
        "start": start.isoformat(),
        "end": end.isoformat(),
        "downloads": dls,
        "uploads": ups,
        "cd": cds,
        "rename": rns,
        "upload_bytes": size,
    }
    try:
        with open(stats_path, "w") as f:
            json.dump(data, f)
        zpath = os.path.join(SESS_DIR, f"session_{sess}.zip")
        with zipfile.ZipFile(zpath, "w") as z:
            z.write(stats_path, os.path.basename(stats_path))
            slog = os.path.join(SESS_DIR, f"{sess}.log")
            if os.path.exists(slog):
                z.write(slog, "session.log")
        summary = (
            f"Session {sess} archived: "
            f"{ups} uploads, {dls} downloads, {rns} renames, {cds} cd"
        )
        alert(
            summary,
            session=sess,
            log_file=os.path.join(SESS_DIR, f"{sess}.log"),
        )
        ACTIONS.extend([
            f"Analyser les logs de {sess}",
            f"Générer un rapport pour {sess}",
            "Nettoyer la quarantine",
        ])
    except Exception:
        pass

def clean_quarantine() -> None:
    """Remove all files from the quarantine directory."""
    removed = 0
    for name in os.listdir(QUAR_DIR):
        path = os.path.join(QUAR_DIR, name)
        try:
            if os.path.isdir(path):
                shutil.rmtree(path)
            else:
                os.remove(path)
            removed += 1
        except Exception as e:
            logging.warning("Failed to remove %s: %s", path, e)
    if removed:
        log_operation(f"CLEAN_QUARANTINE removed {removed} items")
    ACTIONS[:] = [a for a in ACTIONS if not a.startswith("Nettoyer")]
    print(f"{removed} élément(s) supprimé(s) de la quarantine")

# Port-knocking logic
ftp_started = False
knock_state = {}


class AcceptAllChecker:
    """Checker accepting any username/password and returning the username."""
    credentialInterfaces = (credentials.IUsernamePassword,)

    def requestAvatarId(self, creds):
        return defer.succeed(creds.username)
class KnockProtocol(DatagramProtocol):
    def __init__(self, port):
        self.port = port
    def datagramReceived(self, data, addr):
        host = addr[0]
        idx = knock_state.get(host, 0)
        if KNOCK_SEQ[idx] == self.port:
            knock_state[host] = idx + 1
            if knock_state[host] == len(KNOCK_SEQ):
                logging.info("Knock sequence ok from %s", host)
                start_ftp()
                # reset state so additional packets don't raise IndexError
                knock_state[host] = 0
        else:
            knock_state[host] = 0

def start_ftp():
    global ftp_started
    if ftp_started:
        return
    ftp_started = True
    realm = HoneyRealm(ROOT_DIR)
    p     = portal.Portal(realm)
    # Accept any credentials so we can observe attempts
    p.registerChecker(AcceptAllChecker())
    p.registerChecker(AllowAnonymousAccess())
    ctx = ssl.DefaultOpenSSLContextFactory(KEY_FILE, CRT_FILE)
    port_range = range(60000, 60100)
    HoneyFTP.passivePortRange = port_range
    HoneyFTPFactory.passivePortRange = port_range
    factory = HoneyFTPFactory(p, ctx)
    factory.passivePortRange = port_range
    if EXPLICIT:
        HoneyFTP.listenFactory = reactor.listenTCP
        endpoints.TCP4ServerEndpoint(reactor, PORT).listen(factory)
        logging.info("Honeypot FTPES listening on port %s", PORT)
    else:
        def _listen_ssl(port, factory, *a, **kw):
            return reactor.listenSSL(port, factory, ctx, *a, **kw)
        HoneyFTP.listenFactory = staticmethod(_listen_ssl)
        endpoints.SSL4ServerEndpoint(reactor, PORT, ctx).listen(factory)
        logging.info("Honeypot FTPS listening on port %s", PORT)
    try:
        with open(PID_FILE, "w") as f:
            f.write(str(os.getpid()))
    except Exception:
        pass

def randomize_fs(max_dirs=3, max_files=2, max_total=50):
    global dynamic_items
    for p in list(dynamic_items):
        try:
            shutil.rmtree(p) if os.path.isdir(p) else os.remove(p)
        except: pass
        dynamic_items.remove(p)
    try:
        if len(os.listdir(ROOT_DIR)) > max_total:
            return
    except: return

    for _ in range(random.randint(1, max_dirs)):
        try:
            d = os.path.join(ROOT_DIR, f"dir_{uuid.uuid4().hex[:6]}")
            os.makedirs(d, exist_ok=True)
            dynamic_items.append(d)
            for __ in range(random.randint(0, max_files)):
                fpath = os.path.join(d, f"file_{uuid.uuid4().hex[:6]}.txt")
                with open(fpath,"w") as x: x.write("dummy\n")
                dynamic_items.append(fpath)
            sub = os.path.join(d, f"proj_{uuid.uuid4().hex[:4]}")
            os.makedirs(sub, exist_ok=True)
            dynamic_items.append(sub)
            if random.random() > 0.5:
                sf = os.path.join(sub, "README.txt")
                with open(sf, "w") as x: x.write("project notes\n")
                dynamic_items.append(sf)
        except OSError as e:
            logging.warning("randomize_fs OSError: %s", e)
            break

# 7) Custom shell
class HoneyShell(ftp.FTPShell):
    def __init__(self, avatar_id: str):
        super().__init__(filepath.FilePath(ROOT_DIR))
        self.avatarId = avatar_id

    def list(self, path, keys=()):
        d = super().list(path, keys)
        def _inject(res):
            if not path:  # root listing
                res.append(("root.txt", []))
                res.append(("shadow.bak", []))
            return res
        return d.addCallback(_inject)

    def openForReading(self, path):
        rel = "/".join(path)
        if rel in CANARY:
            proto = getattr(self, "protocol", None)
            ip = proto.transport.getPeer().host if proto else None
            sess = proto.session if proto else None
            logf = proto.logf.name if proto and hasattr(proto, "logf") else None
            alert(
                f"CANARY READ {rel}",
                ip=ip,
                user=self.avatarId,
                session=sess,
                log_file=logf,
            )
        abs_path = os.path.join(ROOT_DIR, *path)
        if os.path.isdir(abs_path):
            return defer.fail(ftp.IsADirectoryError(path))
        try:
            f = open(abs_path, "rb")
        except OSError as e:
            return ftp.errnoToFailure(e.errno, path)
        except BaseException:
            return defer.fail()
        else:
            return defer.succeed(ftp._FileReader(f))

    def ftp_CWD(self, path):
        if path.startswith(".."):
            self.logf.write(f"CWD {path}\n")
            log_operation(f"CWD {path} by {self.avatarId}")
            STATS["cd"] += 1
            self.protocol.s_cd += 1 if hasattr(self, 'protocol') else 0
            return ftp.REQ_FILE_ACTN_COMPLETED_OK,
        log_operation(f"CWD {path} by {self.avatarId}")
        STATS["cd"] += 1
        self.protocol.s_cd += 1 if hasattr(self, 'protocol') else 0
        return super().ftp_CWD(path)

# 8) Protocol
class HoneyFTP(ftp.FTP):
    def connectionMade(self):
        super().connectionMade()
        STATS["connections"] += 1
        self.session = get_next_session_id()
        peer        = self.transport.getPeer().host
        self.logf   = open(os.path.join(SESS_DIR,f"{self.session}.log"),"a")
        self.start, self.count = datetime.now(timezone.utc), 0
        self.s_downloads = 0
        self.s_uploads = 0
        self.s_bytes = 0
        self.s_cd = 0
        self.s_ren = 0
        logging.info("CONNECT %s session=%s", peer, self.session)
        if is_tor_exit(peer):
            alert(
                "Tor exit node detected",
                ip=peer,
                session=self.session,
                log_file=self.logf.name,
            )
        self.token = create_honeytoken(peer, self.session)
        self._tls = not EXPLICIT

    def connectionLost(self, reason):
        peer = getattr(self.transport.getPeer(),"host","?")
        logging.info("DISCONNECT %s session=%s", peer, self.session)
        try: os.remove(os.path.join(ROOT_DIR,self.token))
        except: pass
        self.logf.close()
        finalize_session(
            self.session,
            self.start,
            self.s_downloads,
            self.s_uploads,
            self.s_cd,
            self.s_ren,
            self.s_bytes,
        )
        super().connectionLost(reason)

    def ftp_USER(self, u):
        peer = self.transport.getPeer().host
        self.username = u
        logging.info("USER %s %s", peer, u)
        self.logf.write(f"USER {u}\n")
        return super().ftp_USER(u)

    def ftp_PASS(self, pw):
        peer = self.transport.getPeer().host
        attempts = failed_attempts.get(peer, 0)
        if attempts >= BRUTEF_THR:
            delay = DELAY_SEC * (2 ** (attempts - BRUTEF_THR))
            d = defer.Deferred()
            reactor.callLater(
                delay,
                d.callback,
                (ftp.RESPONSE[ftp.AUTH_FAILURE][0],),
            )
            return d
        logging.info("PASS %s %s %s", peer, self.username, pw)
        self.logf.write(f"PASS {pw}\n")
        d = super().ftp_PASS(pw)
        def onFail(e):
            failed_attempts[peer] = failed_attempts.get(peer, 0) + 1
            if failed_attempts[peer] >= BRUTEF_THR:
                alert(
                    "Brute-force attempt",
                    ip=peer,
                    user=self.username,
                    session=self.session,
                    log_file=self.logf.name,
                )
            return e
        def onSucc(r):
            failed_attempts.pop(peer,None)
            randomize_fs()
            lure = create_user_lure(self.username)
            log_operation(f"LOGIN {self.username} session={self.session} lure={lure}")
            STATS["logins"] += 1
            return r
        d.addCallbacks(onSucc,onFail)
        return d

    def ftp_RNFR(self, fn):
        peer = self.transport.getPeer().host
        try:
            old, rel = validate_path(fn, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        self._old = rel
        logging.info("RNFR %s %s", peer, rel)
        with open(os.path.join(SESS_DIR, f"{self.session}.rename.log"), "a") as rl:
            rl.write(f"RNFR {rel}\n")
        log_operation(f"RNFR {old} from {peer} session={self.session}")
        self.sendLine("350 Ready for RNTO")
        return

    def ftp_RNTO(self, new):
        peer = self.transport.getPeer().host
        old_rel = getattr(self, "_old", None)
        try:
            new_path, new_rel = validate_path(new, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        old_path = os.path.join(ROOT_DIR, old_rel) if old_rel else None
        if not old_rel:
            self.sendLine("550 RNFR first")
            return
        src, dst = old_path, new_path
        try:
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            os.rename(src, dst)
            logging.info("RNTO %s %s→%s", peer, old_rel, new_rel)
            with open(os.path.join(SESS_DIR, f"{self.session}.rename.log"), "a") as rl:
                rl.write(f"RNTO {old_rel}→{new_rel}\n")
            log_operation(f"RNTO {old_rel}->{new_rel} from {peer} session={self.session}")
            STATS["renames"] += 1
            self.s_ren += 1
            self.sendLine("250 Rename done")
            return
        except Exception as e:
            self.sendLine(f"550 Rename failed: {e}")
            return

    def ftp_DELE(self, path):
        peer = self.transport.getPeer().host
        try:
            abs_path, rel = validate_path(path, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        tag = f"{self.session}_{uuid.uuid4().hex}"
        dst = os.path.join(QUAR_DIR, tag)
        try:
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            os.replace(abs_path, dst)
            logging.info("DELE %s %s→quarantine/%s", peer, rel, tag)
            self.logf.write(f"DELE {rel}→quarantine/{tag}\n")
            log_operation(f"DELE {rel} by {peer} session={self.session}")
            STATS["deletes"] += 1
            self.sendLine("250 Deleted")
            return
        except Exception as e:
            self.sendLine(f"550 Del failed: {e}")
            return

    def ftp_MKD(self, path):
        try:
            validate_path(path, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        res = ftp.FTP.ftp_MKD(self, path)
        log_operation(f"MKD {path} session={self.session}")
        return res

    def ftp_RMD(self, path):
        try:
            validate_path(path, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        res = ftp.FTP.ftp_RMD(self, path)
        log_operation(f"RMD {path} session={self.session}")
        return res

    def ftp_RETR(self, path):
        peer = self.transport.getPeer().host
        try:
            abs_path, rel = validate_path(path, self.workingDirectory)
            if os.path.isdir(abs_path):
                cache_dir = os.path.join(ROOT_DIR, ".zipcache")
                os.makedirs(cache_dir, exist_ok=True)
                zip_name = f"dir_{uuid.uuid4()}.zip"
                zip_path = os.path.join(cache_dir, zip_name)
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    for root, dirs, files in os.walk(abs_path):
                        for fname in files:
                            full = os.path.join(root, fname)
                            arc = os.path.relpath(full, abs_path)
                            zf.write(full, arc)
                rel_zip = os.path.join(".zipcache", zip_name)
                return super().ftp_RETR(rel_zip)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        if rel in CANARY:
            alert(
                f"CANARY RETR {rel}",
                ip=peer,
                user=getattr(self, "username", None),
                session=self.session,
                log_file=self.logf.name,
            )
        if rel == getattr(self, "token", None):
            logging.info("HONEYTOKEN DL %s session=%s", rel, self.session)
            alert(
                f"HONEYTOKEN {rel}",
                ip=peer,
                user=getattr(self, "username", None),
                session=self.session,
                log_file=self.logf.name,
            )
        self.logf.write(f"RETR {rel}\n")
        log_operation(f"RETR {rel} by {peer} session={self.session}")
        STATS["downloads"] += 1
        self.s_downloads += 1
        return super().ftp_RETR(path)

    def ftp_STOR(self, path):
        peer = self.transport.getPeer().host
        try:
            abs_path, rel = validate_path(path, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        if self.s_bytes >= SESSION_QUOTA or disk_usage(ROOT_DIR) >= QUOTA_BYTES:
            self.sendLine("552 Quota exceeded")
            return
        self.logf.write(f"STOR {rel}\n")
        log_operation(f"STOR {rel} by {peer} session={self.session}")
        STATS["uploads"] += 1
        self.s_uploads += 1
        d = super().ftp_STOR(path)
        def _update(res):
            try:
                self.s_bytes += os.path.getsize(abs_path)
            except Exception:
                pass
            reactor.callInThread(
                sandbox_analyze,
                abs_path,
                self.session,
                getattr(self, "username", "anonymous"),
                peer,
            )
            return res
        d.addCallback(_update)
        return d

    def ftp_NLST(self, path):
        peer = self.transport.getPeer().host
        # When no path is supplied, use an empty string so validate_path()
        # receives a string instead of the working directory list.
        p = path or ""
        try:
            _, rel = validate_path(p, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        self.logf.write(f"NLST {rel}\n")
        log_operation(f"NLST {rel} by {peer} session={self.session}")
        STATS["ls"] += 1
        if "etc" in rel:
            etc = os.path.join(ROOT_DIR, "etc")
            os.makedirs(etc, exist_ok=True)
            with open(os.path.join(etc, "passwd"), "w") as f:
                f.write("root:x:0:0:root:/root:/bin/bash\n")
            with open(os.path.join(etc, "shadow"), "w") as f:
                f.write("root:*:18133:0:99999:7:::\n")
        if "var/log" in rel:
            d = os.path.join(ROOT_DIR, "var/log")
            os.makedirs(d, exist_ok=True)
            with open(os.path.join(d, "syslog.1"), "w") as f:
                f.write("Jan 1 info fake\n")
        return super().ftp_NLST(path)

    def ftp_LIST(self, path=""):
        """Handle LIST even when no path is supplied."""
        if not path:
            # Fallback to NLST for clients that issue bare LIST
            return self.ftp_NLST("")
        return super().ftp_LIST(path)

    def ftp_MLSD(self, path=""):
        """Expose MLSD using our NLST implementation."""
        return self.ftp_NLST(path)

    def ftp_STAT(self, path):
        peer = self.transport.getPeer().host
        try:
            _, rel = validate_path(path, self.workingDirectory)
        except ValueError:
            self.sendLine("550 Invalid path")
            return
        if rel in CANARY:
            alert(
                f"CANARY STAT {rel}",
                ip=peer,
                user=getattr(self, "username", None),
                session=self.session,
                log_file=self.logf.name,
            )
        self.logf.write(f"STAT {rel}\n")
        log_operation(f"STAT {rel} by {peer} session={self.session}")
        return super().ftp_STAT(path)

    def ftp_SITE(self, params):
        parts = params.strip().split(" ",1)
        cmd = parts[0].upper()
        rest = parts[1] if len(parts)>1 else ""
        if cmd=="EXEC" and "/bin/bash" in rest:
            self.sendLine("200 Welcome to Fake Bash v1.0")
            return
        if cmd=="CHMOD":
            self.sendLine("200 CHMOD ignored")
            return
        if cmd=="SHELL":
            self.sendLine("200 SHELL unavailable")
            return
        if cmd=="HELP":
            self.sendLine("200 EXEC CHMOD SHELL DEBUG SQLMAP BOF HELP VERSION GETLOG HISTORY UPTIME STATS")
            return
        if cmd=="VERSION":
            self.sendLine(f"200 {VERSION}")
            return
        if cmd=="GETLOG":
            logfp = OP_LOG if not rest else os.path.join(SESS_DIR, rest+".log")
            try:
                with open(logfp) as f:
                    lines = f.read().splitlines()[-20:]
            except Exception as e:
                lines = [f"log error: {e}"]
            self.sendLine("200-LOG START")
            for l in lines:
                self.sendLine(l)
            self.sendLine("200 LOG END")
            return
        if cmd=="HISTORY":
            target = rest.strip()
            try:
                with open(OP_LOG) as f:
                    lines = [l for l in f.read().splitlines() if target in l][-20:]
            except Exception as e:
                lines = [f"history error: {e}"]
            self.sendLine("200-HISTORY START")
            for l in lines:
                self.sendLine(l)
            self.sendLine("200 HISTORY END")
            return
        if cmd=="DEBUG":
            self.logf.flush()
            try:
                with open(self.logf.name) as f:
                    lines = f.read().splitlines()[-50:]
            except Exception as e:
                lines = [f"log error: {e}"]
            self.sendLine("200-DEBUG START")
            for l in lines:
                self.sendLine(l)
            self.sendLine("200 DEBUG END")
            return
        if cmd=="UPTIME":
            uptime = datetime.now(timezone.utc) - SERVER_START
            self.sendLine(f"200 {str(uptime).split('.')[0]}")
            return
        if cmd=="STATS":
            lines = [
                f"connections={STATS['connections']}",
                f"logins={STATS['logins']}",
                f"uploads={STATS['uploads']}",
                f"downloads={STATS['downloads']}",
                f"deletes={STATS['deletes']}",
                f"renames={STATS['renames']}",
            ]
            self.sendLine("200-STATS START")
            for l in lines:
                self.sendLine(l)
            self.sendLine("200 STATS END")
            return
        if cmd=="SQLMAP":
            self.sendLine("200 SQLi simulation complete")
            return
        if cmd=="BOF":
            if len(rest)>1000:
                logging.warning("SIMUL BOF by %s len=%d", self.transport.getPeer().host, len(rest))
                self.sendLine("451 Buffer overflow!")
                return
            self.sendLine("200")
            return
        # Unknown SITE sub-command → syntax error
        return ftp.SYNTAX_ERR, params

    def ftp_PBSZ(self, param):
        """Handle RFC 4217 PBSZ command used by FTPS clients."""
        self.sendLine("200 PBSZ=0")
        return

    def ftp_PROT(self, param):
        """Acknowledge PROT command without altering connections."""
        level = (param or "").strip().upper()
        if level == "P":
            self.sendLine("200 Protection level set to Private")
        else:
            self.sendLine("200 Protection level ignored")
        return

    def ftp_MODE(self, param):
        p = (param or "").strip().upper()
        if p == "S":
            self.sendLine("200 MODE set to S")
        else:
            self.sendLine("504 Unsupported MODE")
        return

    def ftp_FEAT(self):
        self.sendLine("211-Features:")
        for f in ["UTF8", "PASV", "EPSV", "EPRT", "PBSZ", "PROT", "AUTH TLS"]:
            self.sendLine(" " + f)
        self.sendLine("211 End")
        return

    def ftp_AUTH(self, arg):
        proto = (arg or "").strip().upper()
        if proto in {"TLS", "SSL"}:
            self.sendLine("234 Proceed with negotiation.")
            if EXPLICIT and not getattr(self, "_tls", False):
                self.transport.startTLS(self.factory.ctx)
                self._tls = True
            return
        self.sendLine("504 Unsupported AUTH")
        return

    def ftp_TYPE(self, params):
        """Force binary mode regardless of client request."""
        self.binary = True
        self.sendLine("200 Type set to I")
        return

    def getDTPPort(self, factory, interface=""):
        for portn in self.passivePortRange:
            try:
                if EXPLICIT and not getattr(self, "_tls", False):
                    return reactor.listenTCP(portn, factory, interface=interface)
                return reactor.listenSSL(
                    portn, factory, self.factory.ctx, interface=interface
                )
            except net_error.CannotListenError:
                continue
        raise net_error.CannotListenError("", portn,
                                      f"No port available in range {self.passivePortRange}")

    def lineReceived(self, line):
        peer, cmd = self.transport.getPeer().host, line.decode("latin-1").strip()
        logging.info("CMD %s %s", peer, cmd)
        self.logf.write(cmd+"\n")
        self.count += 1
        now = datetime.now(timezone.utc)
        if self.count > 20 and (now - self.start).total_seconds() < 10:
            alert(
                "Fast scan detected",
                ip=peer,
                session=self.session,
                log_file=self.logf.name,
            )
            self.scan_delay = min(getattr(self, "scan_delay", 0) + 1, 8)
        else:
            self.scan_delay = max(getattr(self, "scan_delay", 0) - 1, 0)
        delay = getattr(self, "scan_delay", 0)
        if delay:
            reactor.callLater(delay, super().lineReceived, line)
        else:
            super().lineReceived(line)
        return

# 9) Factory & Realm
class HoneyFTPFactory(ftp.FTPFactory):
    protocol       = HoneyFTP
    welcomeMessage = random.choice(BANNERS)

    def __init__(self, portal, ctx):
        super().__init__(portal)
        self.ctx = ctx

class HoneyRealm(ftp.FTPRealm):
    def __init__(self, root: str):
        """Use a single virtual filesystem for all users."""
        # FTPRealm expects a plain path string. Passing a FilePath object
        # causes "TypeError: expected str, bytes or os.PathLike object" on
        # startup. Keep the original string here and let FTPRealm convert it.
        super().__init__(root)
        self._root = root

    def requestAvatar(self, avatarId, mind, *interfaces):
        for iface in interfaces:
            if iface is ftp.IFTPShell:
                user = (
                    "anonymous"
                    if avatarId is checkers.ANONYMOUS
                    else str(avatarId)
                )
                return iface, HoneyShell(user), lambda: None
        raise NotImplementedError("Only IFTPShell interface is supported")

# 10) Server thread helpers
server_thread = None
server_running = False
knock_ports = []
reactor_started = False

def run_server():
    """Run the Twisted reactor and bind knock ports once."""
    global knock_ports, reactor_started, server_running
    try:
        if not knock_ports:
            for p in KNOCK_SEQ:
                try:
                    knock_ports.append(reactor.listenUDP(p, KnockProtocol(p)))
                except net_error.CannotListenError:
                    logging.critical(
                        "Port %d déjà utilisé ; modifie KNOCK_SEQ ou libère-le",
                        p,
                    )
                    # Nettoyage : fermer ceux qui étaient déjà ouverts
                    for port in knock_ports:
                        try:
                            port.stopListening()
                        except Exception:
                            pass
                    knock_ports.clear()
                    server_running = False            # on peut retenter plus tard
                    return                            # on sort proprement

        logging.info("Waiting knock sequence %s to start FTP", KNOCK_SEQ)
        if not reactor_started:
            reactor_started = True
            reactor.run(installSignalHandlers=False)
    finally:
        # Si le reactor sort (Ctrl-C ou autre), remettre le flag
        server_running = False

def start_server():
    """Launch the reactor in a background thread if not already running."""
    global server_thread, server_running
    if server_running:
        print("Serveur déjà démarré")
        return
    # Guard: Twisted reactor cannot be restarted once stopped
    if reactor_started:
        print("Reactor déjà arrêté - relance impossible dans ce processus")
        return
    server_running = True
    server_thread = threading.Thread(target=run_server, daemon=True)
    server_thread.start()

def _close_knocks():
    """Close any UDP knock ports currently bound."""
    # Guard: prevents rebinding by clearing the global knock_ports list
    global knock_ports
    for port in knock_ports:
        try:
            port.stopListening()
        except Exception:
            pass
    knock_ports = []

def _shutdown():
    """Stop UDP listeners then terminate the reactor."""
    _close_knocks()
    reactor.stop()

def stop_server():
    """Stop knock listeners and the reactor."""
    global server_running, server_thread
    if not server_running:
        print("Serveur non démarré")
        return
    # Guard: stopListening on ports before halting the reactor
    reactor.callFromThread(_shutdown)
    if server_thread:
        server_thread.join()
        server_thread = None
    server_running = False
    save_stats()
    _cleanup_pid()

def tail_log():
    try:
        with open(LOG_FILE) as f:
            lines = f.read().splitlines()[-20:]
    except FileNotFoundError:
        print("Aucun fichier de log")
        return
    for l in lines:
        print(l)

def list_sessions():
    if not os.path.isdir(SESS_DIR):
        print("Aucune session")
        return
    sess = []
    for f in os.listdir(SESS_DIR):
        if f.endswith('.log') and '.rename' not in f:
            name = f[:-4]
            m = re.search(r'session(\d+)', name)
            idx = int(m.group(1)) if m else 0
            sess.append((idx, name))
    for _, name in sorted(sess):
        print(name)

def list_actions():
    if not ACTIONS:
        print("Aucune action suggérée")
        return
    for idx, action in enumerate(ACTIONS, 1):
        print(f"[{idx}] {action}")


def show_session(sid=None):
    if sid is None:
        sid = input("ID de session > ").strip()
    path = os.path.join(SESS_DIR, f"{sid}.log")
    try:
        with open(path) as f:
            print(f.read())
    except FileNotFoundError:
        print("Session introuvable")

def show_stats():
    for k, v in STATS.items():
        print(f"{k}: {v}")


def generate_report(session_id: str) -> None:
    """Generate a text report for a given session."""
    stats_file = os.path.join(SESS_DIR, f"stats_{session_id}.json")
    log_file = os.path.join(SESS_DIR, f"{session_id}.log")

    try:
        with open(stats_file) as f:
            stats = json.load(f)
    except FileNotFoundError:
        print("Fichier de statistiques introuvable")
        return

    try:
        with open(log_file) as f:
            log_content = f.read()
    except FileNotFoundError:
        log_content = ""

    report_dir = os.path.join(BASE, "reports")
    os.makedirs(report_dir, exist_ok=True)
    report_path = os.path.join(report_dir, f"report_{session_id}.txt")

    with open(report_path, "w") as rep:
        rep.write(f"HoneyFTP Rapport de session {session_id}\n")
        for k, v in stats.items():
            rep.write(f"{k}: {v}\n")
        rep.write("\n")
        rep.write(log_content)

    print(f"Rapport généré: {report_path}")


def menu_loop():
    MENU = (
        "\n"
        "┌───────────────────────────────────────────┐\n"
        "│      High-Interaction FTPS Honeypot       │\n"
        "└───────────────────────────────────────────┘\n"
        "[1] Demarrer le honeypot FTPS\n"
        "[2] Arreter le honeypot\n"
        "[3] Voir les 20 dernieres lignes du log\n"
        "[4] Lister les sessions (ID)\n"
        "[5] Afficher une session (ex. 5)\n"
        "[6] Statistiques globales (connections, uploads...)\n"
        "[7] Actions a effectuer\n"
        "[0] Quitter\n"
    )
    while True:
        print(MENU)
        choice = input("Choix > ").strip()
        if choice == "1":
            start_server()
        elif choice == "2":
            stop_server()
        elif choice == "3":
            tail_log()
        elif choice == "4":
            list_sessions()
        elif choice == "5":
            sid = input("ID de session > ").strip()
            show_session(sid)
        elif choice == "6":
            show_stats()
        elif choice == "7":
            list_actions()
            if ACTIONS:
                sel = input("Numéro de l'action > ").strip()
                try:
                    idx = int(sel) - 1
                    action = ACTIONS[idx]
                except (ValueError, IndexError):
                    print("Choix invalide.")
                    continue
                m = re.search(r"(session\d+)", action)
                sid = m.group(1) if m else None
                if action.startswith("Analyser"):
                    if sid:
                        show_session(sid)
                    else:
                        print("Choix invalide.")
                elif action.startswith("Générer"):
                    if sid:
                        generate_report(sid)
                    else:
                        print("Choix invalide.")
                elif action.startswith("Nettoyer"):
                    clean_quarantine()
                else:
                    print("Choix invalide.")
        elif choice == "0":
            stop_server()
            break


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--server", action="store_true", help=argparse.SUPPRESS)
    args = ap.parse_args()
    if args.server:
        run_server()
    else:
        menu_loop()

if __name__ == "__main__":
    main()
